"""
File Handlers Module for Private Document Q&A System.
Handles text extraction from various document formats.
"""

import os
from pathlib import Path
from typing import Optional, Dict, Any
from loguru import logger

# Document processing imports
import PyPDF2
from docx import Document as DocxDocument

# Local imports
from config.settings import Config


class FileProcessor:
    """
    Handles file processing and text extraction from various document formats.
    
    This class provides functionality to:
    - Extract text from PDF files
    - Extract text from DOCX files
    - Extract text from plain text files
    - Validate file formats
    - Handle file processing errors
    """
    
    def __init__(self, config: Config):
        """
        Initialize the FileProcessor.
        
        Args:
            config: Application configuration
        """
        self.config = config
        self.supported_formats = {
            '.pdf': self._extract_pdf_text,
            '.docx': self._extract_docx_text,
            '.txt': self._extract_txt_text,
            '.md': self._extract_txt_text
        }
    
    def extract_text(self, file_path: Path) -> str:
        """
        Extract text from a document file.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Extracted text content
            
        Raises:
            ValueError: If file format is not supported
            Exception: If text extraction fails
        """
        try:
            # Validate file exists
            if not file_path.exists():
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Get file extension
            file_ext = file_path.suffix.lower()
            
            # Check if format is supported
            if file_ext not in self.supported_formats:
                raise ValueError(f"Unsupported file format: {file_ext}")
            
            # Extract text using appropriate handler
            extractor = self.supported_formats[file_ext]
            text_content = extractor(file_path)
            
            # Validate extracted text
            if not text_content or not text_content.strip():
                raise ValueError(f"No text content extracted from {file_path}")
            
            logger.info(f"Successfully extracted text from {file_path.name}")
            return text_content.strip()
            
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}")
            raise
    
    def _extract_pdf_text(self, file_path: Path) -> str:
        """
        Extract text from a PDF file.
        
        Args:
            file_path: Path to the PDF file
            
        Returns:
            Extracted text content
        """
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # Extract text from each page
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    
                    if page_text:
                        text_content.append(page_text)
            
            return '\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"PDF text extraction failed for {file_path}: {e}")
            raise
    
    def _extract_docx_text(self, file_path: Path) -> str:
        """
        Extract text from a DOCX file.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content
        """
        try:
            doc = DocxDocument(file_path)
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            return '\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"DOCX text extraction failed for {file_path}: {e}")
            raise
    
    def _extract_txt_text(self, file_path: Path) -> str:
        """
        Extract text from a plain text file (TXT, MD).
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Extracted text content
        """
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try with error handling
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                return file.read()
                
        except Exception as e:
            logger.error(f"Text file extraction failed for {file_path}: {e}")
            raise
    
    def get_file_info(self, file_path: Path) -> Dict[str, Any]:
        """
        Get information about a file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Dictionary containing file information
        """
        try:
            stat = file_path.stat()
            
            return {
                'name': file_path.name,
                'size': stat.st_size,
                'extension': file_path.suffix.lower(),
                'is_supported': file_path.suffix.lower() in self.supported_formats,
                'created': stat.st_ctime,
                'modified': stat.st_mtime,
                'path': str(file_path)
            }
            
        except Exception as e:
            logger.error(f"Failed to get file info for {file_path}: {e}")
            return {
                'name': file_path.name,
                'error': str(e)
            }
    
    def validate_file(self, file_path: Path) -> Dict[str, Any]:
        """
        Validate a file for processing.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Validation result dictionary
        """
        result = {
            'valid': True,
            'errors': [],
            'warnings': []
        }
        
        try:
            # Check if file exists
            if not file_path.exists():
                result['valid'] = False
                result['errors'].append("File does not exist")
                return result
            
            # Check file size
            file_size = file_path.stat().st_size
            if file_size > self.config.MAX_FILE_SIZE:
                result['valid'] = False
                result['errors'].append(f"File size ({file_size} bytes) exceeds maximum allowed size")
            
            # Check file extension
            file_ext = file_path.suffix.lower()
            if file_ext not in self.supported_formats:
                result['valid'] = False
                result['errors'].append(f"Unsupported file format: {file_ext}")
            
            # Check if file is readable
            try:
                with open(file_path, 'rb') as f:
                    f.read(1024)  # Try to read first 1KB
            except Exception as e:
                result['valid'] = False
                result['errors'].append(f"Cannot read file: {str(e)}")
            
            # Check for empty files
            if file_size == 0:
                result['warnings'].append("File is empty")
            
            return result
            
        except Exception as e:
            result['valid'] = False
            result['errors'].append(f"Validation error: {str(e)}")
            return result
    
    def get_supported_formats(self) -> Dict[str, str]:
        """
        Get list of supported file formats.
        
        Returns:
            Dictionary mapping file extensions to format descriptions
        """
        return {
            '.pdf': 'Portable Document Format',
            '.docx': 'Microsoft Word Document',
            '.txt': 'Plain Text File',
            '.md': 'Markdown File'
        }
    
    def estimate_processing_time(self, file_path: Path) -> float:
        """
        Estimate processing time for a file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            Estimated processing time in seconds
        """
        try:
            file_size = file_path.stat().st_size
            file_ext = file_path.suffix.lower()
            
            # Base processing rates (bytes per second)
            processing_rates = {
                '.pdf': 10000,  # 10KB/s
                '.docx': 50000,  # 50KB/s
                '.txt': 100000,  # 100KB/s
                '.md': 100000    # 100KB/s
            }
            
            rate = processing_rates.get(file_ext, 10000)
            estimated_time = file_size / rate
            
            # Add minimum processing time
            estimated_time = max(estimated_time, 0.1)
            
            return estimated_time
            
        except Exception:
            return 1.0  # Default estimate 